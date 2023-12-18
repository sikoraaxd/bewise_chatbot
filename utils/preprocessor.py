import os
import re
import tabula
import pandas as pd
from docx import Document as doc

from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.document_loaders.unstructured import UnstructuredFileLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain.docstore.document import Document


def save_files(files, data_path):
    for file in files:
        with open(os.path.join(data_path,file.name), 'wb') as f:
            f.write(file.read())


def clear_files(data_path):
    for file in os.listdir(data_path):
        os.remove(os.path.join(data_path, file))


def parse_tables(filepath):
  extention = filepath.split('.')[-1]
  all_tables = []
  if extention == 'pdf':
      tables = tabula.read_pdf(filepath, pages='all')
      if len(tables):
        for table in tables:
          print(table.to_csv(index=False))
          all_tables.append(
              Document(
                  page_content=table.to_csv(index=False)
              )
          )
  elif extention in ['doc', 'docx']:
    document = doc(filepath)
    tables = []
    for table in document.tables:
        df = [['' for i in range(len(table.columns))] for j in range(len(table.rows))]
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                if cell.text:
                    df[i][j] = cell.text
        tables.append(pd.DataFrame(df))
    if len(tables):
      for table in tables:
        all_tables.append(
            Document(
                page_content=table.to_csv(index=False)
            )
        )
  return all_tables


def parse_documents(data_path):
    documents = []
    for elem in os.listdir(data_path):
        try:
            elem_path = os.path.join(data_path, elem)
            parser = UnstructuredFileLoader(elem_path)
            data = parser.load()
            documents += data
            tables = parse_tables(elem_path)
            documents += tables
        except:
            pass
    return documents


def load_documents(data_path):
    documents = parse_documents(data_path)
    splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
        encoding_name="cl100k_base",
        chunk_size = 1000,
        chunk_overlap = 500,
    )
    documents = splitter.split_documents(documents)
    return documents


def get_vectorstore(documents, openai_api_key):
    vectorstore = FAISS.from_documents(
        documents, embedding=OpenAIEmbeddings(openai_api_key=openai_api_key)
    )
    return vectorstore


def retrieve(query, vectorstore, memory, context, top_k=5):
    retrieved = []
    if vectorstore is not None:
        retrieved = vectorstore.similarity_search(query, k=top_k)
    info_doc = Document(
            page_content="""Используй github markdown по возможности.\nЭти данные тебе дала система: {}\n Дальше идёт основной источник информации\n\n""".
            format(context)
    )
    retrieved = [info_doc, *retrieved]
    if len(memory.buffer_as_messages):
        memory_doc = Document(
            page_content="Дальше следует история диалога:\n{}\n\n".
            format(memory.buffer_as_str)
        )
        retrieved.append(memory_doc)
    print(retrieved[1].page_content)
    return retrieved


def roles_cleaner(text):
    pattern = re.compile(r'^[A-Za-zА-Яа-я.]+[.:]? (.+)$')
    match = pattern.match(text)

    if match:
        return match.group(1)
    else:
        return text


def convert_history_to_memory(history, memory, k=20):
    history = history[-k:]
    i = 0
    while i < len(history)-1:
        if i == 0 and history[i]['role'] == 'assistant':
            i += 1
        _input = history[i]['text']
        _output = history[i+1]['text']
        memory.save_context({'input': _input}, {'output': _output})
        i += 2
    return memory
        
